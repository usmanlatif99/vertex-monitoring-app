import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "scripts" / "create_webauthn_handoff.ps1"
OUT = ROOT / "docs" / "VE-WorkLog-WebAuthn-Attendance-Development-Handoff.docx"

COLORS = {
    "body": "1A2332", "blue": "2E74B5", "dark_blue": "1F4D78",
    "navy": "0B2545", "muted": "5A6A82", "table": "E8EEF5",
}

def rgb(hex_value):
    return RGBColor.from_string(hex_value)

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            mar = tc_pr.find(qn("w:tcMar"))
            if mar is None:
                mar = OxmlElement("w:tcMar")
                tc_pr.append(mar)
            for side, val in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                el = mar.find(qn(f"w:{side}"))
                if el is None:
                    el = OxmlElement(f"w:{side}")
                    mar.append(el)
                el.set(qn("w:w"), str(val))
                el.set(qn("w:type"), "dxa")
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        tr_pr.append(cant_split)

def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar"); fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar"); fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])

def create_numbering_instance(doc):
    style = doc.styles["List Number"]
    base_num_id = style._element.pPr.numPr.numId.val
    numbering = doc.part.numbering_part.element
    base_num = numbering.num_having_numId(base_num_id)
    new_num = numbering.add_num(base_num.abstractNumId.val)
    new_num.add_lvlOverride(ilvl=0).add_startOverride(1)
    return new_num.numId

def apply_num_id(paragraph, num_id):
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = num_id

def style_run(run, size=11, bold=False, color="1A2332", italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)

def format_para(p, before=0, after=6, line=1.25, keep=False):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep

def add_text(doc, text, **kwargs):
    p = doc.add_paragraph()
    style_run(p.add_run(text), kwargs.get("size", 11), kwargs.get("bold", False), kwargs.get("color", COLORS["body"]), kwargs.get("italic", False))
    format_para(p, kwargs.get("before", 0), kwargs.get("after", 6), kwargs.get("line", 1.25), kwargs.get("keep", False))
    return p

def add_heading(doc, text, level):
    spec = {1: (16, COLORS["blue"], 18, 10), 2: (13, COLORS["blue"], 14, 7), 3: (12, COLORS["dark_blue"], 10, 5)}[level]
    p = doc.add_paragraph(style=f"Heading {level}")
    p.clear()
    style_run(p.add_run(text), spec[0], True, spec[1])
    format_para(p, spec[2], spec[3], 1.15, True)
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.first_line_indent = Inches(0)
    if text.startswith("13.") or text.startswith("17."):
        p.paragraph_format.page_break_before = True
    return p

def add_list(doc, text, numbered=False, num_id=None):
    p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
    p.add_run(text)
    if numbered and num_id is not None:
        apply_num_id(p, num_id)
    for run in p.runs: style_run(run)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    format_para(p, 0, 4, 1.25)
    return p

def add_table(doc, headers, rows, widths_in):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, COLORS["table"])
        for run in cell.paragraphs[0].runs: style_run(run, 10.5, True, COLORS["navy"])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                format_para(p, 0, 0, 1.05)
                for run in p.runs: style_run(run, 10)
    widths = [round(v * 1440) for v in widths_in]
    set_table_geometry(table, widths)
    hdr_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader"); repeat.set(qn("w:val"), "true"); hdr_tr_pr.append(repeat)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_callout(doc, label, text, fill):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p1 = cell.paragraphs[0]
    style_run(p1.add_run(label), 10.5, True, COLORS["dark_blue"])
    format_para(p1, 0, 3, 1.15, True)
    p2 = cell.add_paragraph()
    style_run(p2.add_run(text), 10.5)
    format_para(p2, 0, 0, 1.15)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F4F7")
    p = cell.paragraphs[0]
    for idx, line in enumerate(code.strip().splitlines()):
        if idx:
            p.add_run("\n")
        style_run(p.add_run(line), 9, font="Consolas")
    format_para(p, 0, 0, 1.0)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def parse_items(source):
    start = source.index("Add-Table @('Document field'")
    end = source.index("# Save and export PDF")
    body = source[start:end]
    table_re = re.compile(r"Add-Table @\((.*?)\) @\((.*?)\) @\((.*?)\) \| Out-Null", re.S)
    token_re = re.compile(
        r"Add-Table @\(.*?\) @\(.*?\) @\(.*?\) \| Out-Null|"
        r"Add-Callout '([^']*)' '([^']*)' '([^']*)'|"
        r"Add-Heading '([^']*)' ([123])|"
        r"Add-Para '([^']*)'(?: [^\r\n]*)?|"
        r"Add-Bullet '([^']*)'|Add-Number '([^']*)'", re.S
    )
    for m in token_re.finditer(body):
        raw = m.group(0)
        if raw.startswith("Add-Table"):
            tm = table_re.fullmatch(raw)
            headers = re.findall(r"'([^']*)'", tm.group(1))
            rows = [re.findall(r"'([^']*)'", row) for row in re.findall(r"@\((.*?)\)", tm.group(2), re.S)]
            widths = [float(x.strip()) for x in tm.group(3).split(",")]
            yield ("table", headers, rows, widths)
        elif raw.startswith("Add-Callout"):
            yield ("callout", m.group(1), m.group(2), m.group(3))
        elif raw.startswith("Add-Heading"):
            yield ("heading", m.group(4), int(m.group(5)))
        elif raw.startswith("Add-Para"):
            yield ("para", m.group(6))
        elif raw.startswith("Add-Bullet"):
            yield ("bullet", m.group(7))
        else:
            yield ("number", m.group(8))

def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = rgb(COLORS["body"])
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25

    header = sec.header.paragraphs[0]
    style_run(header.add_run("VE WorkLog | WebAuthn Attendance Development Handoff"), 9, False, "6B7A90")
    footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(footer.add_run("Internal technical reference  |  Page "), 9, False, "6B7A90")
    add_page_field(footer)

    add_text(doc, "TECHNICAL DEVELOPMENT HANDOFF", size=10, bold=True, color="1089CC", after=4, keep=True)
    add_text(doc, "WebAuthn-Secured Attendance for VE WorkLog", size=24, bold=True, color=COLORS["navy"], after=6, line=1.05, keep=True)
    add_text(doc, "Device registration, passkey verification, anti-buddy-punching controls, limitations, and implementation specification", size=12, color=COLORS["muted"], after=14, line=1.15, keep=True)

    source = SOURCE.read_text(encoding="utf-8-sig")
    number_group_active = False
    current_num_id = None
    inserted_examples = False
    for item in parse_items(source):
        if item[0] == "heading" and item[1].startswith("7.") and not inserted_examples:
            add_heading(doc, "6.1 Registration options example", 2)
            add_code(doc, """const options = await generateRegistrationOptions({
  rpName: 'VE WorkLog',
  rpID: 'tasks.vertex.pk',
  userName: user.email,
  userDisplayName: user.name,
  attestationType: 'none',
  authenticatorSelection: {
    authenticatorAttachment: 'platform',
    residentKey: 'preferred',
    userVerification: 'required'
  }
});""")
            add_heading(doc, "6.2 Authentication verification example", 2)
            add_code(doc, """const verification = await verifyAuthenticationResponse({
  response: webauthnResponse,
  expectedChallenge: savedChallenge,
  expectedOrigin: 'https://tasks.vertex.pk',
  expectedRPID: 'tasks.vertex.pk',
  credential: storedCredential,
  requireUserVerification: true
});""")
            add_text(doc, "The attendance request should carry latitude, longitude, browser-reported accuracy, and the signed WebAuthn assertion. Bind the saved challenge to the intended action (check-in or checkout) and employee account.")
            inserted_examples = True
        if item[0] == "table": add_table(doc, item[1], item[2], item[3])
        elif item[0] == "callout": add_callout(doc, item[1], item[2], item[3])
        elif item[0] == "heading": add_heading(doc, item[1], item[2])
        elif item[0] == "para": add_text(doc, item[1])
        elif item[0] == "bullet": add_list(doc, item[1])
        elif item[0] == "number":
            if not number_group_active:
                current_num_id = create_numbering_instance(doc)
                number_group_active = True
            add_list(doc, item[1], True, current_num_id)
        if item[0] != "number":
            number_group_active = False

    props = doc.core_properties
    props.title = "WebAuthn-Secured Attendance for VE WorkLog"
    props.subject = "Technical development handoff"
    props.author = "Vertex Electronics"
    doc.save(OUT)
    print(OUT)

if __name__ == "__main__":
    main()
