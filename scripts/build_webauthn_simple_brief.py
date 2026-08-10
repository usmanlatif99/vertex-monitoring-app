from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_webauthn_handoff import (
    COLORS, add_callout, add_heading, add_list, add_page_field, add_table,
    add_text, create_numbering_instance, rgb, style_run,
)

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "VE-WorkLog-WebAuthn-Attendance-Implementation-Brief.docx"

def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = rgb(COLORS["body"])
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.2

    header = sec.header.paragraphs[0]
    style_run(header.add_run("VE WorkLog | WebAuthn Attendance Implementation Brief"), 9, color="6B7A90")
    footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(footer.add_run("Implementation handoff  |  Page "), 9, color="6B7A90")
    add_page_field(footer)

    add_text(doc, "IMPLEMENTATION BRIEF", size=10, bold=True, color="1089CC", after=4, keep=True)
    add_text(doc, "Registered Device + WebAuthn + GPS Attendance", size=23, bold=True, color=COLORS["navy"], after=6, line=1.05, keep=True)
    add_text(doc, "VE WorkLog - Vertex Electronics and Vision Engineering", size=12, color=COLORS["muted"], after=14, keep=True)

    add_callout(doc, "Required result", "An employee can mark normal attendance only when all three checks pass: (1) the employee is logged in, (2) an admin-approved WebAuthn credential completes device user verification, and (3) GPS confirms the device is within the office radius. A password alone must not be enough.", "EAF3F8")

    add_heading(doc, "1. Existing application context", 1)
    add_table(doc, ["Item", "Current value"], [
        ["Application", "Existing Node.js/Express VE WorkLog portal with PostgreSQL and JWT login"],
        ["Production origin", "https://tasks.vertex.pk"],
        ["Office coordinates", "31.441300523433583, 74.32441912480384"],
        ["Office radius", "Approximately 100 metres; validate on the server"],
        ["Attendance status rules", "Before 7:00 AM blocked; after 9:35 AM Late; at/after 11:00 AM Half Day"],
        ["Monthly rules", "3 lates = 1 absence; 2 half-days = 1 absence"],
        ["Testing state", "Employee Attendance navigation is currently hidden; admin Attendance remains available"],
    ], [1.65, 4.85])

    add_heading(doc, "2. Security objective", 1)
    add_text(doc, "The current GPS check proves that a device is at the office but does not prove who is operating it. WebAuthn adds possession of an approved credential and device-level user verification.")
    add_callout(doc, "Attendance decision", "Approved credential + WebAuthn user verification + valid office GPS = accept attendance. If any normal verification step fails, do not silently accept attendance; offer the existing manual request/admin approval process.", "FFF4D6")

    add_heading(doc, "3. Employee device registration", 1)
    registration_num = create_numbering_instance(doc)
    for text in [
        "Employee signs into VE WorkLog on the phone intended for attendance.",
        "Employee opens Attendance and selects Register this device.",
        "Server generates a short-lived, single-use WebAuthn registration challenge.",
        "Browser invokes the platform authenticator and requests fingerprint, Face ID, device PIN, pattern, or another supported verification method.",
        "The private key remains on the device. The public credential response is sent to the server.",
        "Server verifies the challenge, expected origin, relying-party ID, and user-verification result.",
        "Credential is saved as Pending and appears in the admin Device Approvals screen.",
        "Admin approves or rejects it. Policy: one active approved attendance credential per employee.",
    ]:
        add_list(doc, text, numbered=True, num_id=registration_num)

    add_heading(doc, "4. Check-in and checkout", 1)
    attendance_num = create_numbering_instance(doc)
    for text in [
        "Employee selects Check In or Check Out.",
        "Server creates a fresh authentication challenge bound to the employee and requested action.",
        "Browser asks the approved authenticator for user verification and returns a signed assertion.",
        "Browser obtains latitude, longitude, and GPS accuracy.",
        "Frontend sends the signed assertion and GPS payload to the attendance endpoint.",
        "Server verifies WebAuthn first, calculates office distance, checks GPS accuracy, and then applies attendance time/status rules.",
        "Server records attendance plus credential ID, verification time, coordinates, accuracy, IP, and user agent for audit.",
        "Challenge is consumed so it cannot be replayed.",
    ]:
        add_list(doc, text, numbered=True, num_id=attendance_num)

    add_heading(doc, "5. Employee and admin screens", 1)
    add_table(doc, ["Screen/state", "Required behavior"], [
        ["No device", "Show Register this device; block normal check-in/out"],
        ["Pending approval", "Show device and Awaiting admin approval; manual request remains available"],
        ["Approved", "Show normal attendance controls; require WebAuthn for every check-in and checkout"],
        ["Rejected/revoked", "Block normal attendance and offer new/replacement registration"],
        ["Admin Device Approvals", "List employee, device label, browser, registration time, status, Approve/Reject/Revoke"],
        ["Manage Users", "Show device status alongside the existing attendance-enabled toggle"],
    ], [1.75, 4.75])

    add_heading(doc, "6. Proposed database additions", 1)
    add_table(doc, ["Table", "Essential fields"], [
        ["attendance_credentials", "id, user_id, credential_id, public_key, counter, transports, device_name, browser_info, status, registered_at, approved_by, approved_at, revoked_at, revocation_reason"],
        ["webauthn_challenges", "user_id, challenge, purpose/action, expires_at, consumed_at"],
        ["attendance audit additions", "credential reference, verification mode, GPS accuracy, IP address, user agent, verification timestamp"],
    ], [1.85, 4.65])

    add_heading(doc, "7. Required API endpoints", 1)
    add_table(doc, ["Endpoint", "Purpose"], [
        ["POST /api/webauthn/register/options", "Create registration challenge/options"],
        ["POST /api/webauthn/register/verify", "Verify and store pending credential"],
        ["POST /api/webauthn/auth/options", "Create check-in/checkout authentication challenge"],
        ["GET /api/webauthn/device", "Return current employee device status"],
        ["POST /api/webauthn/device/replace", "Request replacement with mandatory reason"],
        ["GET /api/webauthn/admin/devices", "Admin device list"],
        ["PUT /api/webauthn/admin/devices/:id/approve", "Approve pending credential"],
        ["PUT /api/webauthn/admin/devices/:id/reject", "Reject pending credential"],
        ["PUT /api/webauthn/admin/devices/:id/revoke", "Revoke approved credential"],
        ["POST /api/attendance/checkin", "Verify WebAuthn + GPS and record check-in"],
        ["POST /api/attendance/checkout", "Verify WebAuthn + GPS and record checkout"],
    ], [2.85, 3.65])

    add_heading(doc, "8. WebAuthn production settings", 1)
    add_table(doc, ["Setting", "Value"], [
        ["rpName", "VE WorkLog"],
        ["rpID", "tasks.vertex.pk"],
        ["expectedOrigin", "https://tasks.vertex.pk"],
        ["userVerification", "required for registration and authentication"],
        ["authenticatorAttachment", "platform preferred"],
        ["residentKey", "preferred"],
        ["attestationType", "none initially"],
        ["Suggested Node library", "@simplewebauthn/server; follow the installed-version API"],
        ["Challenge policy", "Single-use, action-bound, employee-bound, recommended expiry 3-5 minutes"],
    ], [2.15, 4.35])

    add_heading(doc, "9. Critical limitations", 1)
    add_list(doc, "A website cannot reliably force fingerprint-only or Face-ID-only verification. With userVerification required, the device may allow fingerprint, face, device PIN, pattern, computer password, or security-key PIN.")
    add_list(doc, "VE WorkLog receives only a cryptographic verification result; it does not receive or store biometric data.")
    add_list(doc, "If an employee gives the approved phone and its PIN to a colleague, WebAuthn may still succeed.")
    add_list(doc, "Some passkeys synchronize through cloud credential managers, so a passkey is not always proof of one exact physical phone.")
    add_list(doc, "For stronger protection against deliberate cooperation, add a live selfie/liveness check or use dedicated office biometric hardware.")

    add_heading(doc, "10. Manual fallback and replacement", 1)
    add_list(doc, "If WebAuthn or GPS is unavailable, retain the existing mandatory-reason manual request and admin approval flow.")
    add_list(doc, "Manual records must be labelled unverified by WebAuthn.")
    add_list(doc, "Employee device replacement requires a mandatory reason and admin approval.")
    add_list(doc, "Admin must be able to revoke a lost device immediately; the old credential must stop working after revocation.")
    add_list(doc, "Preserve the existing rule allowing checkout while manual check-in approval is still pending.")

    add_heading(doc, "11. Files expected to change", 1)
    add_table(doc, ["File", "Change"], [
        ["package.json", "Add @simplewebauthn/server; rebuild the Docker image because dependencies change"],
        ["db/schema.sql", "Add credential/challenge tables and attendance audit columns"],
        ["server/routes/webauthn.js", "New registration, authentication, device, and admin routes"],
        ["server/routes/attendance.js", "Require verified assertion for normal attendance and retain manual fallback"],
        ["server/index.js", "Mount /api/webauthn"],
        ["public/app.js", "Device registration/status, WebAuthn ceremony, admin approval UI"],
        ["public/style.css", "Styles for device states and approvals"],
    ], [2.05, 4.45])

    add_heading(doc, "12. Acceptance tests", 1)
    for text in [
        "Correct account password on an unapproved phone cannot mark normal attendance.",
        "Approved credential with failed user verification cannot mark attendance.",
        "Approved credential outside the office radius cannot mark normal attendance.",
        "Expired, reused, wrong-user, or wrong-action challenge is rejected.",
        "Rejected/revoked credential is rejected.",
        "Android Chrome, iPhone Safari, and Windows Hello are tested where supported.",
        "Existing Present/Late/Half Day/absence calculations and daily/monthly/PDF reports still work.",
        "Manual attendance and Early Leave / Remote remarks continue through admin approval.",
    ]:
        add_list(doc, text)

    add_heading(doc, "13. Rollout", 1)
    rollout_num = create_numbering_instance(doc)
    add_list(doc, "Implement device registration and admin approval first without enforcing it.", numbered=True, num_id=rollout_num)
    add_list(doc, "Pilot with a small group across Android and iPhone, plus Windows if desktop attendance is supported.", numbered=True, num_id=rollout_num)
    add_list(doc, "Enable WebAuthn enforcement for pilot employees while retaining manual fallback.", numbered=True, num_id=rollout_num)
    add_list(doc, "Restore Attendance in both employee navigation definitions (empViews and empItems) when testing resumes.", numbered=True, num_id=rollout_num)
    add_list(doc, "After stable operation, enforce it for all attendance-enabled employees.", numbered=True, num_id=rollout_num)
    add_text(doc, "Estimated production-ready implementation: approximately 3-5 working days, including cross-device testing. A basic MVP may be available in about 2 working days but should be piloted before becoming official attendance.")

    add_callout(doc, "Non-negotiable server rule", "Do not trust a client-provided biometric-success flag, office-distance result, or attendance status. Generate and verify challenges on the server, verify origin/RP ID/credential ownership/status, calculate GPS distance on the server, and consume every challenge after use.", "FFF4D6")

    add_heading(doc, "14. Current deployment note", 1)
    add_text(doc, "The application runs in Docker on the Oracle Ubuntu VM (vv_app, vv_db, vv_nginx). Adding a Node package requires rebuilding/recreating the app image; copying JavaScript files into the running container is not enough for new node_modules. The production environment file is /home/ubuntu/vertex-monitoring-app/.env and must remain outside source control.")

    props = doc.core_properties
    props.title = "Registered Device + WebAuthn + GPS Attendance"
    props.subject = "VE WorkLog implementation brief"
    props.author = "Vertex Electronics"
    doc.save(OUT)
    print(OUT)

if __name__ == "__main__":
    main()
